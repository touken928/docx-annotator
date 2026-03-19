"""测试嵌套表格的处理"""

import pytest
from io import BytesIO
from docx import Document as PythonDocxDocument

from docxnote import DocxDocument, Table, Paragraph


class TestNestedTables:
    """测试嵌套表格的读取和批注"""

    def test_nested_table_detection(self, nested_table_doc):
        """测试嵌套表格的检测"""
        dn_doc = DocxDocument.parse(nested_table_doc)

        # 获取外层表格
        outer_tables = [b for b in dn_doc.blocks() if isinstance(b, Table)]
        assert len(outer_tables) > 0

        outer_table = outer_tables[0]
        rows, cols = outer_table.shape()

        # 遍历单元格查找嵌套表格
        found_nested = False
        for r in range(rows):
            for c in range(cols):
                cell = outer_table[r, c]
                for block in cell.blocks():
                    if isinstance(block, Table):
                        found_nested = True
                        break
                if found_nested:
                    break
            if found_nested:
                break

        assert found_nested, "未找到嵌套表格"

    def test_nested_table_structure(self, nested_table_doc):
        """测试嵌套表格的结构"""
        dn_doc = DocxDocument.parse(nested_table_doc)

        outer_tables = [b for b in dn_doc.blocks() if isinstance(b, Table)]
        outer_table = outer_tables[0]

        # 查找包含嵌套表格的单元格
        for r in range(outer_table.shape()[0]):
            for c in range(outer_table.shape()[1]):
                cell = outer_table[r, c]
                for block in cell.blocks():
                    if isinstance(block, Table):
                        # 检查嵌套表格的形状
                        inner_rows, inner_cols = block.shape()
                        assert inner_rows > 0
                        assert inner_cols > 0
                        return

        pytest.fail("未找到嵌套表格")

    def test_nested_table_cell_text(self, nested_table_doc):
        """测试嵌套表格单元格的文本"""
        dn_doc = DocxDocument.parse(nested_table_doc)
        pd_doc = PythonDocxDocument(BytesIO(nested_table_doc))

        # 获取外层表格
        dn_outer = [b for b in dn_doc.blocks() if isinstance(b, Table)][0]
        pd_outer = pd_doc.tables[0]

        # 检查第一个单元格的文本
        dn_cell = dn_outer[0, 0]
        pd_cell = pd_outer.rows[0].cells[0]

        # 提取 docxnote 单元格文本
        dn_texts = []
        for block in dn_cell.blocks():
            if isinstance(block, Paragraph):
                dn_texts.append(block.text)
        dn_text = "\n".join(dn_texts)

        # python-docx 单元格文本
        pd_text = pd_cell.text

        assert dn_text == pd_text

    def test_comment_on_nested_table_cell(self, nested_table_doc):
        """测试为嵌套表格单元格添加批注"""
        dn_doc = DocxDocument.parse(nested_table_doc)

        # 查找嵌套表格
        outer_tables = [b for b in dn_doc.blocks() if isinstance(b, Table)]
        outer_table = outer_tables[0]

        for r in range(outer_table.shape()[0]):
            for c in range(outer_table.shape()[1]):
                cell = outer_table[r, c]
                for block in cell.blocks():
                    if isinstance(block, Table):
                        # 为嵌套表格的单元格添加批注
                        inner_table = block
                        inner_rows, inner_cols = inner_table.shape()

                        if inner_rows > 0 and inner_cols > 0:
                            inner_cell = inner_table[0, 0]
                            for inner_block in inner_cell.blocks():
                                if (
                                    isinstance(inner_block, Paragraph)
                                    and inner_block.text
                                ):
                                    inner_block.comment(
                                        "嵌套表格批注",
                                        end=min(5, len(inner_block.text)),
                                        author="tester",
                                    )
                                    break

                        # 渲染应该成功
                        output = dn_doc.render()
                        assert output is not None
                        assert len(output) > 0
                        return

        pytest.fail("未找到嵌套表格")

    def test_deeply_nested_tables(self):
        """测试深层嵌套表格"""
        # 创建深层嵌套的表格
        pd_doc = PythonDocxDocument()

        # 外层表格
        outer = pd_doc.add_table(rows=1, cols=1)
        outer_cell = outer.rows[0].cells[0]
        outer_cell.text = "外层"

        # 第一层嵌套
        inner1 = outer_cell.add_table(rows=1, cols=1)
        inner1_cell = inner1.rows[0].cells[0]
        inner1_cell.text = "第一层嵌套"

        # 第二层嵌套
        inner2 = inner1_cell.add_table(rows=1, cols=1)
        inner2_cell = inner2.rows[0].cells[0]
        inner2_cell.text = "第二层嵌套"

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)

        # 应该能够访问所有层级
        outer_tables = [b for b in dn_doc.blocks() if isinstance(b, Table)]
        assert len(outer_tables) > 0

        # 遍历查找深层嵌套
        def count_nested_levels(table, level=0):
            max_level = level
            rows, cols = table.shape()
            for r in range(rows):
                for c in range(cols):
                    cell = table[r, c]
                    for block in cell.blocks():
                        if isinstance(block, Table):
                            nested_level = count_nested_levels(block, level + 1)
                            max_level = max(max_level, nested_level)
            return max_level

        max_depth = count_nested_levels(outer_tables[0])
        assert max_depth >= 2, f"嵌套深度应该至少为2，实际为{max_depth}"

    def test_nested_table_with_content(self):
        """测试包含内容的嵌套表格"""
        # 创建包含丰富内容的嵌套表格
        pd_doc = PythonDocxDocument()

        outer = pd_doc.add_table(rows=2, cols=2)

        # 填充外层表格
        outer.rows[0].cells[0].text = "外层 (0,0)"

        # 在 (0,1) 添加嵌套表格
        cell = outer.rows[0].cells[1]
        cell.text = "包含嵌套表格："

        inner = cell.add_table(rows=2, cols=2)
        for i in range(2):
            for j in range(2):
                inner.rows[i].cells[j].text = f"内层 ({i},{j})"

        outer.rows[1].cells[0].text = "外层 (1,0)"
        outer.rows[1].cells[1].text = "外层 (1,1)"

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析并验证
        dn_doc = DocxDocument.parse(doc_bytes)
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))

        # 验证外层表格
        dn_outer = [b for b in dn_doc.blocks() if isinstance(b, Table)][0]
        dn_rows, dn_cols = dn_outer.shape()

        pd_outer = pd_doc2.tables[0]
        pd_rows = len(pd_outer.rows)
        pd_cols = len(pd_outer.columns)

        assert dn_rows == pd_rows
        assert dn_cols == pd_cols

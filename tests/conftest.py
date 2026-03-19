"""pytest 配置和共享 fixtures"""

import pytest
from io import BytesIO
from docx import Document as PythonDocxDocument


@pytest.fixture
def simple_doc():
    """创建简单的测试文档"""
    doc = PythonDocxDocument()

    # 添加标题
    doc.add_heading("测试文档", 0)

    # 添加段落
    doc.add_paragraph("这是第一个段落。")
    doc.add_paragraph("这是第二个段落，包含更多文字用于测试批注功能。")
    doc.add_paragraph("第三个段落包含\n换行符。")

    # 保存到字节流
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@pytest.fixture
def table_doc():
    """创建包含表格的测试文档"""
    doc = PythonDocxDocument()

    doc.add_heading("表格测试文档", 0)
    doc.add_paragraph("下面是一个简单表格：")

    # 创建 3x3 表格
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"

    # 填充表格
    for i in range(3):
        for j in range(3):
            cell = table.rows[i].cells[j]
            cell.text = f"单元格 {i},{j}"

    doc.add_paragraph("表格后的段落。")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@pytest.fixture
def merged_table_doc():
    """创建包含合并单元格的表格文档"""
    doc = PythonDocxDocument()

    doc.add_heading("合并单元格测试", 0)

    # 创建 4x4 表格
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"

    # 填充基础内容
    for i in range(4):
        for j in range(4):
            table.rows[i].cells[j].text = f"R{i}C{j}"

    # 合并第一行的前两个单元格
    cell_a = table.rows[0].cells[0]
    cell_b = table.rows[0].cells[1]
    cell_a.merge(cell_b)
    cell_a.text = "合并单元格 (0,0)-(0,1)"

    # 合并第二列的第2-3行
    cell_c = table.rows[1].cells[1]
    cell_d = table.rows[2].cells[1]
    cell_c.merge(cell_d)
    cell_c.text = "合并单元格 (1,1)-(2,1)"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@pytest.fixture
def nested_table_doc():
    """创建包含嵌套表格的文档"""
    doc = PythonDocxDocument()

    doc.add_heading("嵌套表格测试", 0)
    doc.add_paragraph("外层表格：")

    # 创建外层表格
    outer_table = doc.add_table(rows=2, cols=2)
    outer_table.style = "Table Grid"

    # 第一个单元格：普通文本
    outer_table.rows[0].cells[0].text = "外层单元格 (0,0)"

    # 第二个单元格：包含嵌套表格
    outer_cell = outer_table.rows[0].cells[1]
    outer_cell.text = "外层单元格 (0,1) 包含嵌套表格："

    # 在单元格中添加嵌套表格
    inner_table = outer_cell.add_table(rows=2, cols=2)
    inner_table.style = "Table Grid"
    for i in range(2):
        for j in range(2):
            inner_table.rows[i].cells[j].text = f"内层 {i},{j}"

    # 其他单元格
    outer_table.rows[1].cells[0].text = "外层单元格 (1,0)"
    outer_table.rows[1].cells[1].text = "外层单元格 (1,1)"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@pytest.fixture
def complex_doc():
    """创建复杂的测试文档"""
    doc = PythonDocxDocument()

    # 标题
    doc.add_heading("复杂测试文档", 0)

    # 多个段落
    doc.add_paragraph("第一段：普通文本。")

    p = doc.add_paragraph("第二段：")
    p.add_run("粗体文本").bold = True
    p.add_run(" 和 ")
    p.add_run("斜体文本").italic = True

    doc.add_paragraph("第三段包含特殊字符：<>&\"'")

    # 表格
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for i in range(2):
        for j in range(3):
            table.rows[i].cells[j].text = f"T{i}C{j}"

    # 表格后的段落
    doc.add_paragraph("表格后的段落。")

    # 另一个表格
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = "Table Grid"
    table2.rows[0].cells[0].text = "表格2"
    table2.rows[0].cells[1].text = "数据"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@pytest.fixture
def multiline_cell_doc():
    """创建包含多行文本单元格的文档"""
    doc = PythonDocxDocument()

    doc.add_heading("多行单元格测试", 0)

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    # 单元格包含多个段落
    cell = table.rows[0].cells[0]
    cell.text = "第一行"
    cell.add_paragraph("第二行")
    cell.add_paragraph("第三行")

    table.rows[0].cells[1].text = "单行单元格"
    table.rows[1].cells[0].text = "普通单元格"
    table.rows[1].cells[1].text = "另一个单元格"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

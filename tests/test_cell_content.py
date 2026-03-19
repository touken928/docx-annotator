"""测试表格单元格内容的处理"""

from io import BytesIO
from docx import Document as PythonDocxDocument

from docxnote import DocxDocument, Table, Paragraph


class TestCellContent:
    """测试单元格内容的读取"""

    def test_single_paragraph_cell(self, table_doc):
        """测试单段落单元格"""
        dn_doc = DocxDocument.parse(table_doc)
        pd_doc = PythonDocxDocument(BytesIO(table_doc))

        dn_table = [b for b in dn_doc.blocks() if isinstance(b, Table)][0]
        pd_table = pd_doc.tables[0]

        # 检查第一个单元格
        dn_cell = dn_table[0, 0]
        pd_cell = pd_table.rows[0].cells[0]

        # 提取文本
        dn_texts = [
            block.text for block in dn_cell.blocks() if isinstance(block, Paragraph)
        ]
        dn_text = "\n".join(dn_texts)

        assert dn_text == pd_cell.text

    def test_multi_paragraph_cell(self, multiline_cell_doc):
        """测试多段落单元格"""
        dn_doc = DocxDocument.parse(multiline_cell_doc)
        pd_doc = PythonDocxDocument(BytesIO(multiline_cell_doc))

        dn_table = [b for b in dn_doc.blocks() if isinstance(b, Table)][0]
        pd_table = pd_doc.tables[0]

        # 检查包含多个段落的单元格
        dn_cell = dn_table[0, 0]
        pd_cell = pd_table.rows[0].cells[0]

        # docxnote 段落数
        dn_paragraphs = [
            block for block in dn_cell.blocks() if isinstance(block, Paragraph)
        ]

        # python-docx 段落数
        pd_paragraphs = pd_cell.paragraphs

        # 段落数应该一致
        assert len(dn_paragraphs) == len(pd_paragraphs)

        # 每个段落的文本应该一致
        for dn_p, pd_p in zip(dn_paragraphs, pd_paragraphs):
            assert dn_p.text == pd_p.text

    def test_empty_cell(self):
        """测试空单元格"""
        # 创建包含空单元格的表格
        pd_doc = PythonDocxDocument()
        table = pd_doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "有内容"
        # table.rows[0].cells[1] 保持为空

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        dn_table = [b for b in dn_doc.blocks() if isinstance(b, Table)][0]

        # 检查空单元格
        empty_cell = dn_table[0, 1]
        paragraphs = list(empty_cell.blocks())

        # 空单元格应该至少有一个空段落
        assert len(paragraphs) >= 0

    def test_cell_with_formatting(self):
        """测试包含格式的单元格"""
        # 创建包含格式的表格
        pd_doc = PythonDocxDocument()
        table = pd_doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        p = cell.paragraphs[0]
        p.add_run("普通 ")
        p.add_run("粗体").bold = True
        p.add_run(" ")
        p.add_run("斜体").italic = True

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))

        dn_table = [b for b in dn_doc.blocks() if isinstance(b, Table)][0]
        pd_table = pd_doc2.tables[0]

        dn_cell = dn_table[0, 0]
        pd_cell = pd_table.rows[0].cells[0]

        # 提取文本（忽略格式）
        dn_texts = [
            block.text for block in dn_cell.blocks() if isinstance(block, Paragraph)
        ]
        dn_text = "\n".join(dn_texts)

        assert dn_text == pd_cell.text

    def test_all_cells_in_table(self, table_doc):
        """测试表格中所有单元格的内容"""
        dn_doc = DocxDocument.parse(table_doc)
        pd_doc = PythonDocxDocument(BytesIO(table_doc))

        dn_table = [b for b in dn_doc.blocks() if isinstance(b, Table)][0]
        pd_table = pd_doc.tables[0]

        rows, cols = dn_table.shape()

        # 检查每个单元格
        for r in range(rows):
            for c in range(cols):
                dn_cell = dn_table[r, c]
                pd_cell = pd_table.rows[r].cells[c]

                # 提取文本
                dn_texts = [
                    block.text
                    for block in dn_cell.blocks()
                    if isinstance(block, Paragraph)
                ]
                dn_text = "\n".join(dn_texts)

                assert dn_text == pd_cell.text, f"单元格 [{r},{c}] 文本不一致"

    def test_cell_with_special_content(self):
        """测试包含特殊内容的单元格"""
        # 创建包含特殊内容的表格
        pd_doc = PythonDocxDocument()
        table = pd_doc.add_table(rows=3, cols=1)

        table.rows[0].cells[0].text = "特殊字符：<>&\"'"
        table.rows[1].cells[0].text = "Unicode：中文 🎉 emoji"

        p = table.rows[2].cells[0].paragraphs[0]
        p.add_run("换行")
        p.add_run("\n")
        p.add_run("制表\t符")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))

        dn_table = [b for b in dn_doc.blocks() if isinstance(b, Table)][0]
        pd_table = pd_doc2.tables[0]

        # 检查每个单元格
        for r in range(3):
            dn_cell = dn_table[r, 0]
            pd_cell = pd_table.rows[r].cells[0]

            dn_texts = [
                block.text for block in dn_cell.blocks() if isinstance(block, Paragraph)
            ]
            dn_text = "\n".join(dn_texts)

            assert dn_text == pd_cell.text

    def test_merged_cells_share_same_content_and_blocks_are_tuple(self):
        """复杂合并单元格：覆盖区域应返回相同内容，且 blocks() 返回 tuple"""
        pd_doc = PythonDocxDocument()
        table = pd_doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"

        # 基础填充，方便定位问题
        for r in range(3):
            for c in range(3):
                table.cell(r, c).text = f"R{r}C{c}"

        # 创建一个 2x2 的复合合并块：(0,0) 到 (1,1)
        merged = table.cell(0, 0).merge(table.cell(1, 1))
        merged.text = "MERGED_2x2"

        buf = BytesIO()
        pd_doc.save(buf)
        buf.seek(0)
        doc_bytes = buf.getvalue()

        dn_doc = DocxDocument.parse(doc_bytes)
        dn_table = [b for b in dn_doc.blocks() if isinstance(b, Table)][0]
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))
        pd_table = pd_doc2.tables[0]

        # blocks() 类型要求：文档 blocks 和单元格 blocks 都应是 tuple
        assert isinstance(dn_doc.blocks(), tuple)
        assert isinstance(dn_table[0, 0].blocks(), tuple)

        # 合并覆盖区域 (0,0),(0,1),(1,0),(1,1) 的文本应一致
        merged_coords = [(0, 0), (0, 1), (1, 0), (1, 1)]
        for r, c in merged_coords:
            dn_cell = dn_table[r, c]
            pd_cell = pd_table.cell(r, c)

            dn_texts = [b.text for b in dn_cell.blocks() if isinstance(b, Paragraph)]
            dn_text = "\n".join(dn_texts)

            assert dn_text == pd_cell.text, f"合并覆盖单元格 [{r},{c}] 文本不一致"

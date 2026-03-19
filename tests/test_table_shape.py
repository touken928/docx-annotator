"""测试表格形状理解的一致性"""

from io import BytesIO
from docx import Document as PythonDocxDocument

from docxnote import DocxDocument, Table


class TestTableShape:
    """测试表格形状解析"""

    def test_simple_table_shape(self, table_doc):
        """测试简单表格形状"""
        dn_doc = DocxDocument.parse(table_doc)
        pd_doc = PythonDocxDocument(BytesIO(table_doc))

        dn_tables = [b for b in dn_doc.blocks() if isinstance(b, Table)]
        pd_tables = pd_doc.tables

        assert len(dn_tables) == len(pd_tables)

        for dn_table, pd_table in zip(dn_tables, pd_tables):
            dn_rows, dn_cols = dn_table.shape()
            pd_rows = len(pd_table.rows)
            pd_cols = len(pd_table.columns)

            # 行数应该一致
            assert dn_rows == pd_rows
            # 列数应该一致
            assert dn_cols == pd_cols

    def test_various_table_sizes(self):
        """测试不同尺寸的表格"""
        sizes = [(2, 2), (3, 4), (5, 3), (1, 5), (10, 1)]

        for rows, cols in sizes:
            # 创建表格
            pd_doc = PythonDocxDocument()
            pd_doc.add_table(rows=rows, cols=cols)

            buffer = BytesIO()
            pd_doc.save(buffer)
            buffer.seek(0)
            doc_bytes = buffer.getvalue()

            # 解析
            dn_doc = DocxDocument.parse(doc_bytes)
            pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))

            dn_tables = [b for b in dn_doc.blocks() if isinstance(b, Table)]
            pd_tables = pd_doc2.tables

            dn_rows, dn_cols = dn_tables[0].shape()
            pd_rows = len(pd_tables[0].rows)
            pd_cols = len(pd_tables[0].columns)

            assert dn_rows == pd_rows == rows
            assert dn_cols == pd_cols == cols

    def test_merged_cells_shape(self, merged_table_doc):
        """测试合并单元格的表格形状"""
        dn_doc = DocxDocument.parse(merged_table_doc)
        pd_doc = PythonDocxDocument(BytesIO(merged_table_doc))

        dn_tables = [b for b in dn_doc.blocks() if isinstance(b, Table)]
        pd_tables = pd_doc.tables

        # 表格的整体尺寸应该一致
        dn_rows, dn_cols = dn_tables[0].shape()
        pd_rows = len(pd_tables[0].rows)
        pd_cols = len(pd_tables[0].columns)

        assert dn_rows == pd_rows
        assert dn_cols == pd_cols

    def test_cell_bounds_simple(self, table_doc):
        """测试简单表格的单元格边界"""
        dn_doc = DocxDocument.parse(table_doc)
        dn_tables = [b for b in dn_doc.blocks() if isinstance(b, Table)]

        if dn_tables:
            table = dn_tables[0]
            rows, cols = table.shape()

            # 检查每个单元格的边界
            for r in range(rows):
                for c in range(cols):
                    cell = table[r, c]
                    top, left, bottom, right = cell.bounds()

                    # 对于未合并的单元格
                    assert top == r
                    assert left == c
                    assert bottom == r + 1
                    assert right == c + 1

    def test_cell_bounds_merged(self, merged_table_doc):
        """测试合并单元格的边界"""
        dn_doc = DocxDocument.parse(merged_table_doc)
        dn_tables = [b for b in dn_doc.blocks() if isinstance(b, Table)]

        if dn_tables:
            table = dn_tables[0]

            # 检查合并单元格的边界
            # 第一行前两个单元格合并
            cell_00 = table[0, 0]
            top, left, bottom, right = cell_00.bounds()

            # 边界应该合理
            assert top == 0
            assert left == 0
            assert bottom > top
            assert right > left

    def test_empty_table(self):
        """测试空表格"""
        # 创建空表格（1x1 但无内容）
        pd_doc = PythonDocxDocument()
        pd_doc.add_table(rows=1, cols=1)

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))

        dn_tables = [b for b in dn_doc.blocks() if isinstance(b, Table)]
        pd_tables = pd_doc2.tables

        assert len(dn_tables) == len(pd_tables) == 1

        dn_rows, dn_cols = dn_tables[0].shape()
        assert dn_rows == 1
        assert dn_cols == 1

    def test_multiple_tables(self, complex_doc):
        """测试多个表格"""
        dn_doc = DocxDocument.parse(complex_doc)
        pd_doc = PythonDocxDocument(BytesIO(complex_doc))

        dn_tables = [b for b in dn_doc.blocks() if isinstance(b, Table)]
        pd_tables = pd_doc.tables

        # 表格数量应该一致
        assert len(dn_tables) == len(pd_tables)

        # 每个表格的形状应该一致
        for dn_table, pd_table in zip(dn_tables, pd_tables):
            dn_rows, dn_cols = dn_table.shape()
            pd_rows = len(pd_table.rows)
            pd_cols = len(pd_table.columns)

            assert dn_rows == pd_rows
            assert dn_cols == pd_cols

    def test_merged_cells_bounds_consistency(self):
        """测试合并单元格的 bounds 一致性"""
        # 创建包含合并单元格的表格
        pd_doc = PythonDocxDocument()
        table = pd_doc.add_table(rows=4, cols=4)

        # 填充基础内容
        for i in range(4):
            for j in range(4):
                table.rows[i].cells[j].text = f"R{i}C{j}"

        # 合并第一行的前两个单元格 (0,0) 和 (0,1)
        cell_a = table.rows[0].cells[0]
        cell_b = table.rows[0].cells[1]
        cell_a.merge(cell_b)
        cell_a.text = "合并 (0,0)-(0,1)"

        # 合并第二列的第2-3行 (1,1) 和 (2,1)
        cell_c = table.rows[1].cells[1]
        cell_d = table.rows[2].cells[1]
        cell_c.merge(cell_d)
        cell_c.text = "合并 (1,1)-(2,1)"

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        dn_table = [b for b in dn_doc.blocks() if isinstance(b, Table)][0]

        # 测试水平合并：(0,0) 和 (0,1) 应该有相同的 bounds
        cell_00 = dn_table[0, 0]
        cell_01 = dn_table[0, 1]

        bounds_00 = cell_00.bounds()
        # 访问 bounds_01 以确保实现不抛异常，但不使用返回值
        cell_01.bounds()

        # 根据 API 文档，合并单元格应该返回相同的内容视图
        # 但 bounds 可能不同，因为它们是不同的单元格对象
        # 这里我们验证它们的边界是合理的
        assert bounds_00[0] == 0  # top
        assert bounds_00[1] == 0  # left
        assert bounds_00[2] == 1  # bottom
        assert bounds_00[3] >= 2  # right (至少跨越2列)

        # 测试垂直合并：(1,1) 和 (2,1) 应该有相关的 bounds
        cell_11 = dn_table[1, 1]
        cell_21 = dn_table[2, 1]

        bounds_11 = cell_11.bounds()
        # 访问 bounds_21 以确保实现不抛异常，但不使用返回值
        cell_21.bounds()

        # 验证边界合理性
        assert bounds_11[0] == 1  # top
        assert bounds_11[1] == 1  # left
        assert bounds_11[2] >= 2  # bottom (至少跨越到第2行)
        assert bounds_11[3] == 2  # right

    def test_merged_cells_content_consistency(self):
        """测试合并单元格的内容一致性"""
        # 创建包含合并单元格的表格
        pd_doc = PythonDocxDocument()
        table = pd_doc.add_table(rows=3, cols=3)

        # 合并 (0,0), (0,1), (1,0), (1,1) 成一个大单元格
        cell_00 = table.rows[0].cells[0]
        cell_01 = table.rows[0].cells[1]
        cell_10 = table.rows[1].cells[0]
        cell_11 = table.rows[1].cells[1]

        # 先横向合并
        cell_00.merge(cell_01)
        cell_10.merge(cell_11)
        # 再纵向合并
        cell_00.merge(cell_10)
        cell_00.text = "大合并单元格"

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        dn_table = [b for b in dn_doc.blocks() if isinstance(b, Table)][0]

        # 访问合并区域内的所有单元格
        cells = [
            dn_table[0, 0],
            dn_table[0, 1],
            dn_table[1, 0],
            dn_table[1, 1],
        ]

        # 提取每个单元格的文本
        texts = []
        for cell in cells:
            cell_texts = []
            for block in cell.blocks():
                from docxnote import Paragraph

                if isinstance(block, Paragraph):
                    cell_texts.append(block.text)
            texts.append("\n".join(cell_texts))

        # 根据 API 文档，合并单元格应该返回相同的内容视图
        # 验证至少第一个单元格有内容
        assert len(texts[0]) > 0

        # 验证边界的合理性
        for cell in cells:
            bounds = cell.bounds()
            assert bounds[0] >= 0  # top >= 0
            assert bounds[1] >= 0  # left >= 0
            assert bounds[2] > bounds[0]  # bottom > top
            assert bounds[3] > bounds[1]  # right > left

    def test_complex_merged_cells_bounds(self):
        """测试复杂合并单元格的边界"""
        # 创建复杂的合并场景
        pd_doc = PythonDocxDocument()
        table = pd_doc.add_table(rows=5, cols=5)

        # 填充内容
        for i in range(5):
            for j in range(5):
                table.rows[i].cells[j].text = f"{i},{j}"

        # 多种合并模式
        # 1. 横向合并 (0,0)-(0,2)
        table.rows[0].cells[0].merge(table.rows[0].cells[1])
        table.rows[0].cells[0].merge(table.rows[0].cells[2])

        # 2. 纵向合并 (1,4)-(3,4)
        table.rows[1].cells[4].merge(table.rows[2].cells[4])
        table.rows[1].cells[4].merge(table.rows[3].cells[4])

        # 3. 方块合并 (2,1)-(3,2)
        table.rows[2].cells[1].merge(table.rows[2].cells[2])
        table.rows[3].cells[1].merge(table.rows[3].cells[2])
        table.rows[2].cells[1].merge(table.rows[3].cells[1])

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        dn_table = [b for b in dn_doc.blocks() if isinstance(b, Table)][0]

        # 验证表格尺寸
        rows, cols = dn_table.shape()
        assert rows == 5
        assert cols == 5

        # 验证所有单元格的边界都是合理的
        for r in range(rows):
            for c in range(cols):
                cell = dn_table[r, c]
                top, left, bottom, right = cell.bounds()

                # 基本合理性检查
                assert 0 <= top < rows
                assert 0 <= left < cols
                assert top < bottom <= rows
                assert left < right <= cols

                # 边界应该包含当前单元格
                assert top <= r < bottom
                assert left <= c < right

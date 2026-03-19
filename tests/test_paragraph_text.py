"""测试段落文本理解的一致性"""

from io import BytesIO
from docx import Document as PythonDocxDocument

from docxnote import DocxDocument, Paragraph


class TestParagraphText:
    """测试段落文本提取与 python-docx 的一致性"""

    def test_simple_text(self, simple_doc):
        """测试简单文本提取"""
        dn_doc = DocxDocument.parse(simple_doc)
        pd_doc = PythonDocxDocument(BytesIO(simple_doc))

        dn_paragraphs = [b for b in dn_doc.blocks() if isinstance(b, Paragraph)]
        pd_paragraphs = pd_doc.paragraphs

        for dn_p, pd_p in zip(dn_paragraphs, pd_paragraphs):
            assert dn_p.text == pd_p.text

    def test_text_with_formatting(self):
        """测试带格式的文本提取"""
        # 创建带格式的文档
        pd_doc = PythonDocxDocument()
        p = pd_doc.add_paragraph()
        p.add_run("普通文本 ")
        p.add_run("粗体文本").bold = True
        p.add_run(" ")
        p.add_run("斜体文本").italic = True

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))

        dn_paragraphs = [b for b in dn_doc.blocks() if isinstance(b, Paragraph)]
        pd_paragraphs = pd_doc2.paragraphs

        # 文本内容应该一致（忽略格式）
        assert dn_paragraphs[0].text == pd_paragraphs[0].text

    def test_text_with_line_breaks(self):
        """测试包含换行符的文本"""
        # 创建包含换行的文档
        pd_doc = PythonDocxDocument()
        p = pd_doc.add_paragraph()
        p.add_run("第一行")
        p.add_run("\n")
        p.add_run("第二行")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))

        dn_paragraphs = [b for b in dn_doc.blocks() if isinstance(b, Paragraph)]
        pd_paragraphs = pd_doc2.paragraphs

        # 文本应该包含换行符
        assert dn_paragraphs[0].text == pd_paragraphs[0].text
        assert "\n" in dn_paragraphs[0].text

    def test_text_with_tabs(self):
        """测试包含制表符的文本"""
        # 创建包含制表符的文档
        pd_doc = PythonDocxDocument()
        p = pd_doc.add_paragraph()
        p.add_run("列1")
        p.add_run("\t")
        p.add_run("列2")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))

        dn_paragraphs = [b for b in dn_doc.blocks() if isinstance(b, Paragraph)]
        pd_paragraphs = pd_doc2.paragraphs

        # 文本应该包含制表符
        assert dn_paragraphs[0].text == pd_paragraphs[0].text
        assert "\t" in dn_paragraphs[0].text

    def test_special_characters(self):
        """测试特殊字符"""
        # 创建包含特殊字符的文档
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("特殊字符：<>&\"'")
        pd_doc.add_paragraph("Unicode：中文 🎉 emoji")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))

        dn_paragraphs = [b for b in dn_doc.blocks() if isinstance(b, Paragraph)]
        pd_paragraphs = pd_doc2.paragraphs

        for dn_p, pd_p in zip(dn_paragraphs, pd_paragraphs):
            assert dn_p.text == pd_p.text

    def test_empty_paragraph(self):
        """测试空段落"""
        # 创建空段落
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("")
        pd_doc.add_paragraph("非空段落")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))

        dn_paragraphs = [b for b in dn_doc.blocks() if isinstance(b, Paragraph)]
        pd_paragraphs = pd_doc2.paragraphs

        assert dn_paragraphs[0].text == ""
        assert pd_paragraphs[0].text == ""

    def test_whitespace_preservation(self):
        """测试空白字符保留"""
        # 创建包含多个空格的文档
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("多个    空格")
        pd_doc.add_paragraph("  前导空格")
        pd_doc.add_paragraph("尾随空格  ")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 解析
        dn_doc = DocxDocument.parse(doc_bytes)
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))

        dn_paragraphs = [b for b in dn_doc.blocks() if isinstance(b, Paragraph)]
        pd_paragraphs = pd_doc2.paragraphs

        for dn_p, pd_p in zip(dn_paragraphs, pd_paragraphs):
            assert dn_p.text == pd_p.text

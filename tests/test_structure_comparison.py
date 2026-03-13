"""测试文档结构与 python-docx 的一致性"""

import pytest
from io import BytesIO
from docx import Document as PythonDocxDocument

from docxnote import DocxDocument, Paragraph, Table


class TestStructureComparison:
    """测试文档结构解析的一致性"""
    
    def test_paragraph_order(self, simple_doc):
        """测试段落顺序与 python-docx 一致"""
        # docxnote
        dn_doc = DocxDocument.parse(simple_doc)
        dn_paragraphs = [b for b in dn_doc.blocks() if isinstance(b, Paragraph)]
        
        # python-docx
        pd_doc = PythonDocxDocument(BytesIO(simple_doc))
        pd_paragraphs = pd_doc.paragraphs
        
        # 段落数量应该一致
        assert len(dn_paragraphs) == len(pd_paragraphs)
        
        # 段落文本应该一致
        for dn_p, pd_p in zip(dn_paragraphs, pd_paragraphs):
            assert dn_p.text == pd_p.text
    
    def test_table_order(self, complex_doc):
        """测试表格顺序与 python-docx 一致"""
        # docxnote
        dn_doc = DocxDocument.parse(complex_doc)
        dn_blocks = list(dn_doc.blocks())
        
        # python-docx
        pd_doc = PythonDocxDocument(BytesIO(complex_doc))
        
        # 提取 docxnote 的表格
        dn_tables = [b for b in dn_blocks if isinstance(b, Table)]
        
        # python-docx 的表格
        pd_tables = pd_doc.tables
        
        # 表格数量应该一致
        assert len(dn_tables) == len(pd_tables)
    
    def test_mixed_blocks_order(self, complex_doc):
        """测试段落和表格混合顺序"""
        # docxnote
        dn_doc = DocxDocument.parse(complex_doc)
        dn_blocks = list(dn_doc.blocks())
        
        # python-docx
        pd_doc = PythonDocxDocument(BytesIO(complex_doc))
        
        # 记录 docxnote 的块类型序列
        dn_sequence = []
        for block in dn_blocks:
            if isinstance(block, Paragraph):
                dn_sequence.append("P")
            elif isinstance(block, Table):
                dn_sequence.append("T")
        
        # 记录 python-docx 的块类型序列
        # 注意：python-docx 没有统一的块迭代器，需要手动构建
        pd_sequence = []
        
        # 获取所有段落和表格的位置
        para_indices = {id(p): i for i, p in enumerate(pd_doc.paragraphs)}
        table_indices = {id(t): i for i, t in enumerate(pd_doc.tables)}
        
        # 简化：只检查表格数量和段落数量
        dn_para_count = dn_sequence.count("P")
        dn_table_count = dn_sequence.count("T")
        
        assert dn_para_count == len(pd_doc.paragraphs)
        assert dn_table_count == len(pd_doc.tables)
    
    def test_empty_paragraphs_preserved(self, simple_doc):
        """测试空段落是否被保留"""
        # 创建包含空段落的文档
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("第一段")
        pd_doc.add_paragraph("")  # 空段落
        pd_doc.add_paragraph("第三段")
        
        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()
        
        # docxnote
        dn_doc = DocxDocument.parse(doc_bytes)
        dn_paragraphs = [b for b in dn_doc.blocks() if isinstance(b, Paragraph)]
        
        # python-docx
        pd_doc2 = PythonDocxDocument(BytesIO(doc_bytes))
        pd_paragraphs = pd_doc2.paragraphs
        
        # 数量应该一致
        assert len(dn_paragraphs) == len(pd_paragraphs)
        
        # 检查空段落
        assert dn_paragraphs[1].text == ""
        assert pd_paragraphs[1].text == ""

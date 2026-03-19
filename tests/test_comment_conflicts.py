"""测试批注冲突和边界情况"""

import zipfile
from io import BytesIO
from lxml import etree
from docx import Document as PythonDocxDocument

from docxnote import DocxDocument, Paragraph


class TestCommentConflicts:
    """测试批注的 Run 冲突和重叠情况"""

    def test_overlapping_comments(self):
        """测试重叠批注"""
        # 创建包含文本 "1234" 的文档
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("1234")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 第一次批注 "23"
        doc = DocxDocument.parse(doc_bytes)
        paragraphs = [b for b in doc.blocks() if isinstance(b, Paragraph)]
        paragraphs[0].comment("批注23", start=1, end=3, author="user1")

        # 渲染
        output1 = doc.render()

        # 验证第一次渲染的 XML 合法性
        with zipfile.ZipFile(BytesIO(output1)) as z:
            doc_xml = z.read("word/document.xml")
            doc_tree = etree.fromstring(doc_xml)
            assert doc_tree is not None

            comments_xml = z.read("word/comments.xml")
            comments_tree = etree.fromstring(comments_xml)
            assert len(comments_tree) == 1

        # 第二次解析并批注 "12"
        doc2 = DocxDocument.parse(output1, keep_comments=True)
        paragraphs2 = [b for b in doc2.blocks() if isinstance(b, Paragraph)]
        paragraphs2[0].comment("批注12", start=0, end=2, author="user2")

        # 再次渲染
        output2 = doc2.render()

        # 验证第二次渲染的 XML 合法性
        with zipfile.ZipFile(BytesIO(output2)) as z:
            doc_xml = z.read("word/document.xml")
            doc_tree = etree.fromstring(doc_xml)
            assert doc_tree is not None

            comments_xml = z.read("word/comments.xml")
            comments_tree = etree.fromstring(comments_xml)
            # 应该有两个批注
            assert len(comments_tree) == 2

        # 验证文本内容没有改变
        doc3 = DocxDocument.parse(output2, keep_comments=True)
        paragraphs3 = [b for b in doc3.blocks() if isinstance(b, Paragraph)]
        assert paragraphs3[0].text == "1234"

    def test_adjacent_comments(self):
        """测试相邻批注"""
        # 创建文档
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("ABCDEF")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 添加相邻批注
        doc = DocxDocument.parse(doc_bytes)
        paragraphs = [b for b in doc.blocks() if isinstance(b, Paragraph)]

        # "AB"
        paragraphs[0].comment("批注AB", start=0, end=2, author="user1")
        # "CD"
        paragraphs[0].comment("批注CD", start=2, end=4, author="user2")
        # "EF"
        paragraphs[0].comment("批注EF", start=4, end=6, author="user3")

        output = doc.render()

        # 验证 XML 合法性
        with zipfile.ZipFile(BytesIO(output)) as z:
            doc_xml = z.read("word/document.xml")
            doc_tree = etree.fromstring(doc_xml)
            assert doc_tree is not None

            comments_xml = z.read("word/comments.xml")
            comments_tree = etree.fromstring(comments_xml)
            assert len(comments_tree) == 3

        # 验证文本内容
        doc2 = DocxDocument.parse(output, keep_comments=True)
        paragraphs2 = [b for b in doc2.blocks() if isinstance(b, Paragraph)]
        assert paragraphs2[0].text == "ABCDEF"

    def test_nested_comments_same_position(self):
        """测试同一位置的嵌套批注"""
        # 创建文档
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("12345678")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 添加嵌套批注
        doc = DocxDocument.parse(doc_bytes)
        paragraphs = [b for b in doc.blocks() if isinstance(b, Paragraph)]

        # 外层批注 "1234"
        paragraphs[0].comment("外层批注", start=0, end=4, author="user1")
        # 内层批注 "23"
        paragraphs[0].comment("内层批注", start=1, end=3, author="user2")

        output = doc.render()

        # 验证 XML 合法性
        with zipfile.ZipFile(BytesIO(output)) as z:
            doc_xml = z.read("word/document.xml")
            doc_tree = etree.fromstring(doc_xml)
            assert doc_tree is not None

            comments_xml = z.read("word/comments.xml")
            comments_tree = etree.fromstring(comments_xml)
            assert len(comments_tree) == 2

    def test_multiple_render_cycles(self):
        """测试多次渲染循环"""
        # 创建文档
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("ABCDEFGHIJ")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 第一轮：批注 "ABC"
        doc1 = DocxDocument.parse(doc_bytes)
        p1 = [b for b in doc1.blocks() if isinstance(b, Paragraph)][0]
        p1.comment("批注1", start=0, end=3, author="user1")
        output1 = doc1.render()

        # 第二轮：批注 "DEF"
        doc2 = DocxDocument.parse(output1, keep_comments=True)
        p2 = [b for b in doc2.blocks() if isinstance(b, Paragraph)][0]
        p2.comment("批注2", start=3, end=6, author="user2")
        output2 = doc2.render()

        # 第三轮：批注 "GHI"
        doc3 = DocxDocument.parse(output2, keep_comments=True)
        p3 = [b for b in doc3.blocks() if isinstance(b, Paragraph)][0]
        p3.comment("批注3", start=6, end=9, author="user3")
        output3 = doc3.render()

        # 验证最终结果
        with zipfile.ZipFile(BytesIO(output3)) as z:
            doc_xml = z.read("word/document.xml")
            doc_tree = etree.fromstring(doc_xml)
            assert doc_tree is not None

            comments_xml = z.read("word/comments.xml")
            comments_tree = etree.fromstring(comments_xml)
            assert len(comments_tree) == 3

        # 验证文本内容
        doc_final = DocxDocument.parse(output3, keep_comments=True)
        p_final = [b for b in doc_final.blocks() if isinstance(b, Paragraph)][0]
        assert p_final.text == "ABCDEFGHIJ"

    def test_comment_on_already_commented_text(self):
        """测试对已有批注的文本再次批注"""
        # 创建文档
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("1234567890")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 第一次批注 "345"
        doc1 = DocxDocument.parse(doc_bytes)
        p1 = [b for b in doc1.blocks() if isinstance(b, Paragraph)][0]
        p1.comment("第一次批注", start=2, end=5, author="user1")
        output1 = doc1.render()

        # 验证第一次批注
        with zipfile.ZipFile(BytesIO(output1)) as z:
            comments_xml = z.read("word/comments.xml")
            comments_tree = etree.fromstring(comments_xml)
            assert len(comments_tree) == 1

        # 第二次批注 "234"（与第一次重叠）
        doc2 = DocxDocument.parse(output1, keep_comments=True)
        p2 = [b for b in doc2.blocks() if isinstance(b, Paragraph)][0]
        p2.comment("第二次批注", start=1, end=4, author="user2")
        output2 = doc2.render()

        # 验证第二次批注
        with zipfile.ZipFile(BytesIO(output2)) as z:
            doc_xml = z.read("word/document.xml")
            doc_tree = etree.fromstring(doc_xml)
            assert doc_tree is not None

            comments_xml = z.read("word/comments.xml")
            comments_tree = etree.fromstring(comments_xml)
            assert len(comments_tree) == 2

        # 第三次批注 "456"（与前两次都重叠）
        doc3 = DocxDocument.parse(output2, keep_comments=True)
        p3 = [b for b in doc3.blocks() if isinstance(b, Paragraph)][0]
        p3.comment("第三次批注", start=3, end=6, author="user3")
        output3 = doc3.render()

        # 验证第三次批注
        with zipfile.ZipFile(BytesIO(output3)) as z:
            doc_xml = z.read("word/document.xml")
            doc_tree = etree.fromstring(doc_xml)
            assert doc_tree is not None

            comments_xml = z.read("word/comments.xml")
            comments_tree = etree.fromstring(comments_xml)
            assert len(comments_tree) == 3

        # 验证文本内容始终不变
        doc_final = DocxDocument.parse(output3, keep_comments=True)
        p_final = [b for b in doc_final.blocks() if isinstance(b, Paragraph)][0]
        assert p_final.text == "1234567890"

    def test_full_text_comment_then_partial(self):
        """测试先批注全文，再批注部分"""
        # 创建文档
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("ABCDE")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        # 第一次批注全文
        doc1 = DocxDocument.parse(doc_bytes)
        p1 = [b for b in doc1.blocks() if isinstance(b, Paragraph)][0]
        p1.comment("全文批注", start=0, end=5, author="user1")
        output1 = doc1.render()

        # 第二次批注部分 "BC"
        doc2 = DocxDocument.parse(output1, keep_comments=True)
        p2 = [b for b in doc2.blocks() if isinstance(b, Paragraph)][0]
        p2.comment("部分批注", start=1, end=3, author="user2")
        output2 = doc2.render()

        # 验证 XML 合法性
        with zipfile.ZipFile(BytesIO(output2)) as z:
            doc_xml = z.read("word/document.xml")
            doc_tree = etree.fromstring(doc_xml)
            assert doc_tree is not None

            comments_xml = z.read("word/comments.xml")
            comments_tree = etree.fromstring(comments_xml)
            assert len(comments_tree) == 2

        # 验证文本内容
        doc3 = DocxDocument.parse(output2, keep_comments=True)
        p3 = [b for b in doc3.blocks() if isinstance(b, Paragraph)][0]
        assert p3.text == "ABCDE"

    def test_comment_boundaries_edge_cases(self):
        """测试批注边界的边缘情况"""
        # 创建文档
        pd_doc = PythonDocxDocument()
        pd_doc.add_paragraph("0123456789")

        buffer = BytesIO()
        pd_doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()

        doc = DocxDocument.parse(doc_bytes)
        p = [b for b in doc.blocks() if isinstance(b, Paragraph)][0]

        # 边界情况：起始位置相同
        p.comment("批注1", start=0, end=3, author="user1")
        p.comment("批注2", start=0, end=5, author="user2")

        # 边界情况：结束位置相同
        p.comment("批注3", start=5, end=10, author="user3")
        p.comment("批注4", start=7, end=10, author="user4")

        # 边界情况：完全包含
        p.comment("批注5", start=2, end=8, author="user5")

        output = doc.render()

        # 验证 XML 合法性
        with zipfile.ZipFile(BytesIO(output)) as z:
            doc_xml = z.read("word/document.xml")
            doc_tree = etree.fromstring(doc_xml)
            assert doc_tree is not None

            comments_xml = z.read("word/comments.xml")
            comments_tree = etree.fromstring(comments_xml)
            assert len(comments_tree) == 5

        # 验证文本内容
        doc2 = DocxDocument.parse(output, keep_comments=True)
        p2 = [b for b in doc2.blocks() if isinstance(b, Paragraph)][0]
        assert p2.text == "0123456789"
